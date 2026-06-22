import pandas as pd
import io
import os
import re


# ─── Shared rubric definitions (used by evaluation form, detail view, PDFs) ──

# Supervisor rubric: 5 questions x /10 = /50
SUPERVISOR_QUESTIONS = [
    ('score_clarity', {
        'title': 'Clarity of problem & significance of study',
        'description': 'Clarity of problem being investigated and significance of study'
                       ' (Score from 1 to 10, where 10 is the highest mark and 1 is the lowest mark)'
    }),
    ('score_literature', {
        'title': 'Literature knowledge & acknowledgment',
        'description': "Candidate's knowledge of pertinent literature and correct acknowledgment"
                       " without plagiarizing (Score from 1 to 10, where 10 is the highest mark"
                       " and 1 is the lowest mark)"
    }),
    ('score_knowledge', {
        'title': 'Subject knowledge, methodology & conclusions',
        'description': 'Evidence of sound knowledge and understanding of the subject studied,'
                       ' suitability of techniques, clarity and justification of conclusions'
                       ' and/or recommendations (Score from 1 to 10, where 10 is the highest mark'
                       ' and 1 is the lowest mark)'
    }),
    ('score_critical', {
        'title': 'Critical analysis & original contribution',
        'description': 'Critical discussion and sound analysis of results obtained, relevance'
                       ' and importance of findings and original contribution to knowledge'
                       ' by way of examples or results (Score from 1 to 10, where 10 is the highest'
                       ' mark and 1 is the lowest mark)'
    }),
    ('score_presentation', {
        'title': 'Presentation quality',
        'description': 'Presentation (language, format, standard of presentation, sectioning,'
                       ' grammar, typing, citation and referencing, etc.) (Score from 1 to 10,'
                       ' where 10 is the highest mark and 1 is the lowest mark)'
    }),
]

# Examiner / Jury President rubric: 8 questions = /100
EXAMINER_QUESTIONS = [
    ('score_intro', {
        'title': 'Introduction of topic',
        'description': 'Score from 1 to 5, where 5 is the highest mark and 1 is the lowest mark',
        'min': 1, 'max': 5,
    }),
    ('score_transition', {
        'title': 'Smooth transition between sections (Logical flow)',
        'description': 'Score from 1 to 10, where 10 is the highest mark and 1 is the lowest mark',
        'min': 1, 'max': 10,
    }),
    ('score_quality', {
        'title': 'Quality of content (Scientific and technical content)',
        'description': 'Score from 1 to 10, where 10 is the highest mark and 1 is the lowest mark',
        'min': 1, 'max': 10,
    }),
    ('score_analysis', {
        'title': 'Student analysis (analytical content)',
        'description': 'Score from 1 to 5, where 5 is the highest mark and 1 is the lowest mark',
        'min': 1, 'max': 5,
    }),
    ('score_correctness', {
        'title': 'Clarity and correctness of content',
        'description': 'Score from 1 to 5, where 5 is the highest mark and 1 is the lowest mark',
        'min': 1, 'max': 5,
    }),
    ('score_summary', {
        'title': 'Summary and conclusions',
        'description': 'Score from 1 to 5, where 5 is the highest mark and 1 is the lowest mark',
        'min': 1, 'max': 5,
    }),
]

# Presentation block a-e: each scored 2,4,6,8,10 (step 2, max 10) -> total /50
PRESENTATION_BLOCK = [
    ('score_delivery',   'a- Professional and confident delivery'),
    ('score_engagement', 'b- Engaged with audience (holding audience\u2019s attention)'),
    ('score_structure',  'c- Overall structure and organization of presentation'),
    ('score_timing',     'd- Length of talk (delivered in assigned time)'),
    ('score_qa',         'e- Response to questions from audience'),
]

# Final question for Examiner/President: 1-10
OVERALL_QUESTION = ('score_overall', {
    'title': 'Overall Impression/Quality',
    'description': 'Score from 1 to 10, where 10 is the highest mark and 1 is the lowest mark',
    'min': 1, 'max': 10,
})

EVALUATION_ROLES = [
    ('supervisor', 'Supervisor'),
    ('examiner', 'Examiner'),
    ('president', 'Jury President'),
]


def clean_name(val):
    if pd.isna(val):
        return ''
    return str(val).strip().replace('\n', ' ')


def import_excel_students(session, file_obj):
    from .models import Student
    df = pd.read_excel(file_obj, engine='openpyxl')
    count = 0
    for _, row in df.iterrows():
        surname = clean_name(row.get('Surname', ''))
        given = clean_name(row.get('Given Names', ''))
        if not surname and not given:
            continue

        room_raw = row.get('Room ', row.get('Room', 1))
        try:
            room = int(str(room_raw).strip())
        except:
            room = 1

        def_date = row.get('Date', None)
        if def_date is not None and pd.isna(def_date):
            def_date = None
        elif def_date is not None and hasattr(def_date, 'date'):
            def_date = def_date.date()

        s = Student(
            session=session,
            time=clean_name(row.get('Time (GMT)', '')),
            surname=surname,
            given_names=given,
            thesis_title=clean_name(row.get('Research Project title', '')),
            room=room,
            defense_date=def_date,
            supervisor_1_name=clean_name(row.get('Academic Supervisors', '')),
            supervisor_1_email=clean_name(row.get('Email of Supervisor ', row.get('Email of Supervisor', ''))),
            supervisor_1_institution=clean_name(row.get('Institution of Supervisors', '')),
            supervisor_2_name=clean_name(row.get('Academic Supervisors 2', '')),
            supervisor_2_email=clean_name(row.get('Email of Supervisor 2', '')),
            jury_president_name=clean_name(row.get('President Jury', '')),
            jury_president_email=clean_name(row.get('Email', '')),
            jury_president_affiliation=clean_name(row.get('Affiliation ', row.get('Affiliation', ''))),
            examiner_1_name=clean_name(row.get('Examiner 1', '')),
            examiner_1_email=clean_name(row.get('Examiner 1 email', '')),
            examiner_1_affiliation=clean_name(row.get('Affiliation', '')),
            examiner_2_name=clean_name(row.get('Examiner 2', '')),
            examiner_2_email=clean_name(row.get('Examiner 2 email', '')),
        )
        s.save()
        count += 1
    return count


def calc_scores(student, evaluations):
    supervisor_evals = [e for e in evaluations if e.role == 'supervisor']
    examiner_evals = [e for e in evaluations if e.role in ('examiner', 'president')]

    sup_scores = [e.total_score_percent() for e in supervisor_evals]
    exam_scores = [e.total_score_percent() for e in examiner_evals]

    sup_avg = sum(sup_scores) / len(sup_scores) if sup_scores else 0
    exam_avg = sum(exam_scores) / len(exam_scores) if exam_scores else 0
    final = 0.4 * sup_avg + 0.6 * exam_avg if (sup_scores or exam_scores) else 0

    def grade(pct):
        if pct >= 85: return "Distinction"
        elif pct >= 80: return "Very Good Pass"
        elif pct >= 70: return "Good Pass"
        elif pct >= 60: return "Pass"
        else: return "Fail"

    return {
        'supervisor_evals': supervisor_evals,
        'examiner_evals': examiner_evals,
        'sup_avg': round(sup_avg, 2),
        'exam_avg': round(exam_avg, 2),
        'final_score': round(final, 2),
        'final_grade': grade(final),
    }


def _get_logo_path():
    """
    Return absolute path to the header logo.
    Uses the centralized PVSettings.logo if set, otherwise falls back
    to the default AIMS logo bundled with the app.
    """
    try:
        from .models import PVSettings
        settings_obj = PVSettings.get_solo()
        if settings_obj.logo and os.path.exists(settings_obj.logo.path):
            return settings_obj.logo.path
    except Exception:
        pass
    base = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base, 'static', 'defense', 'img', 'aims-logo-header.png')


def _get_pv_signature():
    """Return (director_name, director_title) from centralized PVSettings."""
    try:
        from .models import PVSettings
        settings_obj = PVSettings.get_solo()
        return settings_obj.director_name, settings_obj.director_title
    except Exception:
        return "Dr. Coura Balde", "Academic Manager, AIMS Senegal"


def generate_pv_docx(student, evaluations):
    """
    Generates a DOCX matching the exact format of PV_Defense-Abass_Ndiaye.docx:
    - Header: AIMS logo (centered, same dimensions as original)
    - US Letter page, 1-inch margins
    - Arial font throughout
    - Bold+underline centred title "THESIS DEFENSE REPORT"
    - Fields: Defense date, Student, Supervisor(s), Dissertation title
    - Two-column table: Committee members | Decision  (with grey final row)
    - GRADE CATEGORY section (underlined heading, small font list)
    - Right-aligned signature block: Dr. Coura Balde / Academic Manager, AIMS Senegal
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    import copy

    scores = calc_scores(student, evaluations)

    doc = Document()

    # ── Page: US Letter (21.59cm x 27.94cm), 2.54cm (1in) margins ─────────────
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)
        section.header_distance = Inches(0.49)

    # ── Header: AIMS logo ─────────────────────────────────────────────────────
    logo_path = _get_logo_path()
    header = doc.sections[0].header
    # Clear default paragraph
    for p in header.paragraphs:
        p.clear()
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run()
    if os.path.exists(logo_path):
        # Original image dimensions in original: cx=3057816 EMU, cy=560928 EMU
        # ≈ 3.34 inches wide × 0.61 inches tall
        hr.add_picture(logo_path, width=Inches(3.34), height=Inches(0.61))

    # ── Helper: Arial run ─────────────────────────────────────────────────────
    def set_arial(run, size_pt=12, bold=False, underline=False, color=None):
        run.font.name = 'Arial'
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.underline = underline
        if color:
            run.font.color.rgb = RGBColor(*color)
        # force Arial for all scripts
        rpr = run._r.get_or_add_rPr()
        for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
            el = OxmlElement(f'w:rFonts')
            el.set(qn('w:' + attr), 'Arial')
            existing = rpr.find(qn('w:rFonts'))
            if existing is None:
                rpr.insert(0, el)
            else:
                existing.set(qn('w:' + attr), 'Arial')

    def add_para(alignment=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = alignment
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        return p

    # ── 3 blank lines (matching original spacing above title) ─────────────────
    for _ in range(3):
        add_para()

    # ── Title: THESIS DEFENSE REPORT ─────────────────────────────────────────
    title_p = add_para(WD_ALIGN_PARAGRAPH.CENTER)
    title_r = title_p.add_run('THESIS DEFENSE REPORT')
    set_arial(title_r, size_pt=16, bold=True, underline=True)

    # blank line
    add_para()
    add_para()

    # ── Defense date / Student / Supervisor / Dissertation title ──────────────
    # Reference: each field is bold+underlined, 12pt, with NO blank paragraphs
    # between them (zero additional spacing).
    def add_field(label, value):
        p = add_para()
        r1 = p.add_run(label)
        set_arial(r1, size_pt=12, bold=True, underline=False)
        r2 = p.add_run(' ' + str(value))
        set_arial(r2, size_pt=12, bold=False, underline=False)

    defense_date_str = ''
    if student.defense_date:
        d = student.defense_date
        if isinstance(d, str):
            defense_date_str = d
        else:
            defense_date_str = d.strftime('%B %d, %Y')

    add_field('Defense date:', defense_date_str)

    # Student line: "Given Names Surname - ID: <academic_year>"
    student_id = f"{student.session.academic_year}"
    add_field('Student:', f'{student.full_name()} - ID: {student_id}')

    # Supervisor(s)
    add_field('Supervisor:', student.supervisor_1_name)
    if student.supervisor_2_name:
        add_field('Co-Supervisor:', student.supervisor_2_name)

    # Dissertation title - strip any URLs, wraps naturally onto next line
    clean_title = re.sub(r'https?://\S+', '', student.thesis_title).strip()
    clean_title = re.sub(r'\s+', ' ', clean_title).rstrip(':').strip()
    add_field('Dissertation title:', clean_title)

    add_para()

    # ── Committee table ───────────────────────────────────────────────────────
    # Col widths from original: 6866 + 2451 = 9317 DXA
    # 1 DXA = 1/20 pt = 1/1440 inch  →  6866/1440=4.77in, 2451/1440=1.70in
    COL1 = Inches(4.77)
    COL2 = Inches(1.70)
    TABLE_W = Inches(6.47)

    def make_border(val='single', sz=4, color='000000', space=0):
        b = OxmlElement('w:' + 'top')  # placeholder, caller replaces tag
        b.set(qn('w:val'), val)
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:color'), color)
        b.set(qn('w:space'), str(space))
        return b

    def set_cell_border(cell, top=True, bottom=True, left=True, right=True,
                        top_nil=False, left_nil=False):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')

        def border_el(tag, nil=False):
            el = OxmlElement(f'w:{tag}')
            if nil:
                el.set(qn('w:val'), 'nil')
                el.set(qn('w:sz'), '0')
                el.set(qn('w:color'), '000000')
                el.set(qn('w:space'), '0')
            else:
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '4')
                el.set(qn('w:color'), '000000')
                el.set(qn('w:space'), '0')
            return el

        tcBorders.append(border_el('top',    nil=top_nil))
        tcBorders.append(border_el('left',   nil=left_nil))
        tcBorders.append(border_el('bottom', nil=not bottom))
        tcBorders.append(border_el('right',  nil=not right))
        tcPr.append(tcBorders)

        # vertical align center
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)

    def set_cell_shading(cell, fill='auto'):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Set table width
    tblPr = table._tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9317')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    # Table borders
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        b.set(qn('w:space'), '0')
        tblBorders.append(b)
    tblPr.append(tblBorders)

    # Grid
    tblGrid = OxmlElement('w:tblGrid')
    for w in (6866, 2451):
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    table._tbl.insert(table._tbl.index(tblPr) + 1, tblGrid)

    def add_table_row(col1_text, col2_text, header=False, grey=False,
                      col1_size=11.5, col2_size=12, first_row=False, bold=None):
        if bold is None:
            bold = header or grey
        row = table.add_row()
        c1, c2 = row.cells[0], row.cells[1]

        # widths
        for cell, w in ((c1, 6866), (c2, 2451)):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(w))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

        # borders
        set_cell_border(c1, top_nil=not first_row, left_nil=False)
        set_cell_border(c2, top_nil=not first_row, left_nil=not first_row)

        # shading
        fill = 'd9d9d9' if grey else 'auto'
        set_cell_shading(c1, fill)
        set_cell_shading(c2, fill)

        # content col 1
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after  = Pt(0)
        r1 = p1.add_run(col1_text)
        set_arial(r1, size_pt=col1_size, bold=bold, underline=False)

        # content col 2
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)
        r2 = p2.add_run(col2_text)
        set_arial(r2, size_pt=col2_size, bold=bold, underline=False)

    # Header row
    add_table_row('Committee members', 'Decision', header=True, first_row=True)

    # Data rows – each evaluator
    for ev in evaluations:
        name_with_affil = ev.evaluator_name
        # append institution if available
        if ev.role == 'supervisor' and student.supervisor_1_institution:
            if ev.evaluator_name == student.supervisor_1_name:
                name_with_affil = f"{ev.evaluator_name}, {student.supervisor_1_institution}"
        elif ev.role == 'president' and student.jury_president_affiliation:
            name_with_affil = f"{ev.evaluator_name}, {student.jury_president_affiliation}"
        elif ev.role == 'examiner':
            if ev.evaluator_name == student.examiner_1_name and student.examiner_1_affiliation:
                name_with_affil = f"{ev.evaluator_name}, {student.examiner_1_affiliation}"
            elif ev.evaluator_name == student.examiner_2_name and student.examiner_2_affiliation:
                name_with_affil = f"{ev.evaluator_name}, {student.examiner_2_affiliation}"
        add_table_row(name_with_affil, ev.grade_label(), col1_size=11.5)

    # Build final grade formula correctly
    examiners  = [e for e in evaluations if e.role == 'examiner']
    presidents = [e for e in evaluations if e.role == 'president']
    n_ex = len(examiners)
    n_pr = len(presidents)
    n_total = n_ex + n_pr   # denominator for the 60% portion

    parts = [f'Examiner {i+1}' for i in range(n_ex)]
    parts += ['Jury President'] * n_pr
    if n_total == 0:
        final_formula = 'Final Grade = 40%(Supervisor) + 60%(Examiner/1)'
    elif n_total == 1:
        final_formula = f'Final Grade = 40%(Supervisor) + 60%({parts[0]})'
    else:
        inner = ' + '.join(parts)
        final_formula = f'Final Grade = 40%(Supervisor) + 60%(({inner})/{n_total})'

    add_table_row(final_formula, scores['final_grade'], grey=True, col1_size=13.5, col2_size=14)

    add_para()

    # ── GRADE CATEGORY ────────────────────────────────────────────────────────
    p_gc = add_para()
    r_gc = p_gc.add_run('GRADE CATEGORY')
    set_arial(r_gc, size_pt=9, bold=True, underline=True)

    for cat in ['Fail (0-59%)', 'Pass (60-69%)', 'Good Pass (70-79%)',
                'Very Good Pass (80-85%)', 'Distinction (85+%)']:
        p_cat = add_para()
        r_cat = p_cat.add_run(cat)
        set_arial(r_cat, size_pt=7.5)

    # spacers
    for _ in range(4):
        add_para()

    # ── Signature block ─────────────────────────────────────────────────────
    # Positioned on the right side of the page, but the two lines are
    # left-aligned WITHIN that block so "Dr." sits directly above "Academic"
    # (D under A) instead of each line being right-aligned independently.
    director_name, director_title = _get_pv_signature()

    sig_table = doc.add_table(rows=1, cols=2)
    sig_tbl_pr = sig_table._tbl.tblPr
    sig_tblW = OxmlElement('w:tblW')
    sig_tblW.set(qn('w:w'), '9317')
    sig_tblW.set(qn('w:type'), 'dxa')
    sig_tbl_pr.append(sig_tblW)

    # No borders
    sig_borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'nil')
        sig_borders.append(b)
    sig_tbl_pr.append(sig_borders)

    sig_grid = OxmlElement('w:tblGrid')
    for w in (5817, 3500):  # left spacer column, right signature column
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        sig_grid.append(gc)
    sig_table._tbl.insert(sig_table._tbl.index(sig_tbl_pr) + 1, sig_grid)

    sig_cell_left, sig_cell_right = sig_table.rows[0].cells

    for cell, width in ((sig_cell_left, 5817), (sig_cell_right, 3500)):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(width))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    # Left cell stays empty (acts as a spacer to push signature right)
    sig_cell_left.paragraphs[0].paragraph_format.space_before = Pt(0)
    sig_cell_left.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Right cell: two left-aligned lines, "Dr." sits above "Academic"
    p_sig1 = sig_cell_right.paragraphs[0]
    p_sig1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sig1.paragraph_format.space_before = Pt(0)
    p_sig1.paragraph_format.space_after = Pt(0)
    r_sig1 = p_sig1.add_run(director_name)
    set_arial(r_sig1, size_pt=10, bold=True, underline=False)

    p_sig2 = sig_cell_right.add_paragraph()
    p_sig2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sig2.paragraph_format.space_before = Pt(0)
    p_sig2.paragraph_format.space_after = Pt(0)
    r_sig2 = p_sig2.add_run(director_title)
    set_arial(r_sig2, size_pt=10, bold=False, underline=False)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_pv_pdf(student, evaluations):
    """PDF version matching the same layout as the DOCX."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, Image as RLImage, HRFlowable)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    scores = calc_scores(student, evaluations)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        topMargin=inch, bottomMargin=inch,
        leftMargin=inch, rightMargin=inch
    )

    BLACK  = colors.black
    GREY   = colors.HexColor('#D9D9D9')
    story  = []

    # ── Logo ──────────────────────────────────────────────────────────────────
    logo_path = _get_logo_path()
    if os.path.exists(logo_path):
        logo = RLImage(logo_path, width=3.34*inch, height=0.61*inch)
        logo.hAlign = 'CENTER'
        story.append(logo)
    story.append(Spacer(1, 0.3*inch))

    # ── Styles ────────────────────────────────────────────────────────────────
    FONT = 'Helvetica'
    FONT_B = 'Helvetica-Bold'

    title_style  = ParagraphStyle('Title',  fontName=FONT_B, fontSize=16,
                                  alignment=TA_CENTER, underlineProportion=0.1,
                                  spaceAfter=4)
    normal_style = ParagraphStyle('Normal', fontName=FONT, fontSize=12, spaceAfter=0, spaceBefore=0, leading=14)
    small_style  = ParagraphStyle('Small',  fontName=FONT, fontSize=7.5, spaceAfter=1, leading=9)
    small_b_style= ParagraphStyle('SmallB', fontName=FONT_B, fontSize=9,
                                  underlineProportion=0.1, spaceAfter=2)
    left_b       = ParagraphStyle('LB', fontName=FONT_B, fontSize=10,
                                  alignment=TA_LEFT, spaceAfter=1)
    left_n       = ParagraphStyle('LN', fontName=FONT,  fontSize=10,
                                  alignment=TA_LEFT)
    final_cell_style = ParagraphStyle('FinalCell', fontName=FONT_B, fontSize=12, leading=14)
    final_cell_style_center = ParagraphStyle('FinalCellC', fontName=FONT_B, fontSize=12, leading=14, alignment=TA_CENTER)

    # ── Title ─────────────────────────────────────────────────────────────────
    story.append(Paragraph('<u><b>THESIS DEFENSE REPORT</b></u>', title_style))
    story.append(Spacer(1, 0.15*inch))

    def field(label, value):
        story.append(Paragraph(f'<b>{label}</b> {value}', normal_style))

    defense_date_str = ''
    if student.defense_date:
        d = student.defense_date
        defense_date_str = d.strftime('%B %d, %Y') if hasattr(d, 'strftime') else str(d)

    field('Defense date:', defense_date_str)
    field('Student:', f'{student.full_name()} - ID: {student.session.academic_year}')
    field('Supervisor:', student.supervisor_1_name)
    if student.supervisor_2_name:
        field('Co-Supervisor:', student.supervisor_2_name)

    clean_title = re.sub(r'https?://\S+', '', student.thesis_title).strip()
    clean_title = re.sub(r'\s+', ' ', clean_title).rstrip(':').strip()
    field('Dissertation title:', clean_title)
    story.append(Spacer(1, 0.1*inch))

    # ── Committee table ───────────────────────────────────────────────────────
    col1_w = 4.77 * inch
    col2_w = 1.70 * inch

    border = ('GRID', (0, 0), (-1, -1), 0.5, BLACK)

    cell_style = ParagraphStyle('Cell', fontName=FONT, fontSize=10, leading=12)
    cell_style_center = ParagraphStyle('CellC', fontName=FONT, fontSize=10, leading=12, alignment=TA_CENTER)
    header_cell_style = ParagraphStyle('HeaderCell', fontName=FONT_B, fontSize=10, leading=12)
    header_cell_style_center = ParagraphStyle('HeaderCellC', fontName=FONT_B, fontSize=10, leading=12, alignment=TA_CENTER)

    t_data = [[Paragraph('Committee members', header_cell_style),
               Paragraph('Decision', header_cell_style_center)]]
    for ev in evaluations:
        name_with_affil = ev.evaluator_name
        if ev.role == 'supervisor' and student.supervisor_1_institution:
            if ev.evaluator_name == student.supervisor_1_name:
                name_with_affil = f"{ev.evaluator_name}, {student.supervisor_1_institution}"
        elif ev.role == 'president' and student.jury_president_affiliation:
            name_with_affil = f"{ev.evaluator_name}, {student.jury_president_affiliation}"
        elif ev.role == 'examiner':
            if ev.evaluator_name == student.examiner_1_name and student.examiner_1_affiliation:
                name_with_affil = f"{ev.evaluator_name}, {student.examiner_1_affiliation}"
            elif ev.evaluator_name == student.examiner_2_name and student.examiner_2_affiliation:
                name_with_affil = f"{ev.evaluator_name}, {student.examiner_2_affiliation}"
        t_data.append([
            Paragraph(name_with_affil, cell_style),
            Paragraph(ev.grade_label(), cell_style_center),
        ])

    examiners  = [e for e in evaluations if e.role == 'examiner']
    presidents = [e for e in evaluations if e.role == 'president']
    n_ex = len(examiners)
    n_pr = len(presidents)
    n_total = n_ex + n_pr

    parts = [f'Examiner {i+1}' for i in range(n_ex)]
    parts += ['Jury President'] * n_pr
    if n_total == 0:
        final_formula = 'Final Grade = 40%(Supervisor) + 60%(Examiner/1)'
    elif n_total == 1:
        final_formula = f'Final Grade = 40%(Supervisor) + 60%({parts[0]})'
    else:
        inner = ' + '.join(parts)
        final_formula = f'Final Grade = 40%(Supervisor) + 60%(({inner})/{n_total})'

    t_data.append([
        Paragraph(final_formula, final_cell_style),
        Paragraph(scores['final_grade'], final_cell_style_center),
    ])

    last = len(t_data) - 1
    tbl = Table(t_data, colWidths=[col1_w, col2_w])
    tbl.setStyle(TableStyle([
        border,
        ('VALIGN',    (0, 0),    (-1, -1),    'MIDDLE'),
        ('BACKGROUND',(0, last), (-1, last),  GREY),
        ('TOPPADDING',(0, 0),    (-1, -1),    4),
        ('BOTTOMPADDING',(0,0),  (-1, -1),    4),
        ('LEFTPADDING', (0,0),   (-1, -1),    4),
        ('RIGHTPADDING',(0,0),   (-1, -1),    4),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 0.15*inch))

    # ── Grade categories ──────────────────────────────────────────────────────
    story.append(Paragraph('<u><b>GRADE CATEGORY</b></u>', small_b_style))
    for cat in ['Fail (0-59%)', 'Pass (60-69%)', 'Good Pass (70-79%)',
                'Very Good Pass (80-85%)', 'Distinction (85+%)']:
        story.append(Paragraph(cat, small_style))

    story.append(Spacer(1, 0.6*inch))

    # ── Signature ─────────────────────────────────────────────────────────────
    # Positioned on the right side of the page, with both lines left-aligned
    # within their cell so "Dr." sits directly above "Academic" (D under A).
    director_name, director_title = _get_pv_signature()
    sig_inner = [
        [Paragraph(f'<b>{director_name}</b>', left_b)],
        [Paragraph(director_title, left_n)],
    ]
    sig_table = Table(sig_inner, colWidths=[2.2*inch])
    sig_table.setStyle(TableStyle([
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))

    outer = Table([['', sig_table]], colWidths=[4.3*inch, 2.2*inch])
    outer.setStyle(TableStyle([
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    story.append(outer)

    doc.build(story)
    buf.seek(0)
    return buf


def get_evaluation_detail_rows(evaluation):
    """
    Build a list of (question_label, answer_text) rows for a single
    evaluation, covering every question that evaluator answered -
    used by both the on-screen Detail page and the Detail PDF.
    """
    rows = []
    rows.append(('Adresse e-mail', evaluation.evaluator_email))
    rows.append(('Title/Full name', evaluation.evaluator_name))
    rows.append(('Student full name', evaluation.student_full_name))
    rows.append(('Title of project', evaluation.thesis_title_confirmed))
    rows.append(('You are', evaluation.get_role_display()))
    rows.append(('Research paper coming out from the project',
                  'Yes' if evaluation.research_paper else 'No'))

    if evaluation.role == 'supervisor':
        for field_name, question in SUPERVISOR_QUESTIONS:
            score = getattr(evaluation, field_name, 0)
            rows.append((question['title'], f'{score} / 10'))
        rows.append(('Section Total', f'{evaluation.total_score_raw()} / 50'))
    else:
        for field_name, question in EXAMINER_QUESTIONS:
            score = getattr(evaluation, field_name, 0)
            rows.append((question['title'], f'{score} / {question["max"]}'))

        for field_name, label in PRESENTATION_BLOCK:
            score = getattr(evaluation, field_name, 0)
            rows.append((label, f'{score} / 10'))
        rows.append(('Presentation Subtotal', f'{evaluation.presentation_block_raw()} / 50'))

        ov_field, ov_question = OVERALL_QUESTION
        score = getattr(evaluation, ov_field, 0)
        rows.append((ov_question['title'], f'{score} / {ov_question["max"]}'))

        rows.append(('Total Score', f'{evaluation.total_score_raw()} / 100'))

    rows.append(('Grade', evaluation.grade_label()))
    rows.append(('Comments', evaluation.additional_info or '-'))
    rows.append(('Annotated document uploaded',
                  'Yes' if evaluation.annotated_document else 'No'))
    return rows


def generate_detail_pdf(student, evaluations):
    """
    Generate a PDF containing the full breakdown of every submitted
    evaluation for this student - every question and the answer given,
    for Supervisor(s), Examiner(s) and Jury President.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, Image as RLImage, PageBreak)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        topMargin=inch, bottomMargin=inch,
        leftMargin=inch, rightMargin=inch
    )

    FONT = 'Helvetica'
    FONT_B = 'Helvetica-Bold'
    AIMS_RED = colors.HexColor('#6B0B1F')
    GREY = colors.HexColor('#F0F0F0')

    title_style = ParagraphStyle('DTitle', fontName=FONT_B, fontSize=15,
                                  alignment=TA_CENTER, textColor=AIMS_RED, spaceAfter=10)
    h_style = ParagraphStyle('DH', fontName=FONT_B, fontSize=11,
                              textColor=colors.white, spaceAfter=0)
    normal = ParagraphStyle('DN', fontName=FONT, fontSize=9, leading=12)
    normal_b = ParagraphStyle('DNB', fontName=FONT_B, fontSize=9, leading=12)
    info = ParagraphStyle('Info', fontName=FONT, fontSize=10, spaceAfter=2)

    story = []

    logo_path = _get_logo_path()
    if os.path.exists(logo_path):
        logo = RLImage(logo_path, width=3.34*inch, height=0.61*inch)
        logo.hAlign = 'CENTER'
        story.append(logo)
        story.append(Spacer(1, 0.2*inch))

    story.append(Paragraph('EVALUATION DETAIL REPORT', title_style))
    story.append(Paragraph(f'<b>Student:</b> {student.full_name()}', info))
    story.append(Paragraph(f'<b>Thesis title:</b> {student.thesis_title}', info))
    story.append(Paragraph(f'<b>Defense date:</b> {student.defense_date or ""}', info))
    story.append(Spacer(1, 0.2*inch))

    if not evaluations:
        story.append(Paragraph('No evaluations have been submitted for this student.', normal))

    for idx, ev in enumerate(evaluations):
        if idx > 0:
            story.append(PageBreak())

        header_text = f'{ev.evaluator_name} - {ev.get_role_display()}'
        header_table = Table([[Paragraph(header_text, h_style)]], colWidths=[6.5*inch])
        header_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), AIMS_RED),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(header_table)
        story.append(Spacer(1, 0.1*inch))

        rows = get_evaluation_detail_rows(ev)
        table_data = []
        for label, value in rows:
            table_data.append([
                Paragraph(str(label), normal_b),
                Paragraph(str(value), normal),
            ])

        t = Table(table_data, colWidths=[3.5*inch, 3*inch])
        style = [
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ]
        # Highlight total/grade rows
        for i, (label, _) in enumerate(rows):
            if 'Total' in label or label == 'Grade' or 'Subtotal' in label:
                style.append(('BACKGROUND', (0, i), (-1, i), GREY))
        t.setStyle(TableStyle(style))
        story.append(t)

    doc.build(story)
    buf.seek(0)
    return buf
